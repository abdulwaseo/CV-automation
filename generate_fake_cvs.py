from faker import Faker
import random
from docx import Document
import os
import shutil

fake = Faker()
output_dir = "generated_cvs"
os.makedirs(output_dir, exist_ok=True)

skills_pool = ["Communication", "Teamwork", "Leadership"]

def generate_cv(index):
    doc = Document()
    name = fake.name()
    email = fake.email()
    selected_skills = ", ".join(random.sample(skills_pool, k=random.randint(0, 2)))
    education = f"{random.choice(['BSc', 'BS', 'MSc'])} in {random.choice(['Computer Science', 'Software Engineering', 'Information Technology'])} from {fake.company()} University"
    experience = f"{random.randint(1, 10)}+ years experience in backend and cloud systems."

    doc.add_heading(name, 0)
    doc.add_paragraph(f"Email: {email}")
    doc.add_paragraph(f"\nSkills:\n{selected_skills}")
    doc.add_paragraph(f"\nExperience:\n{experience}")
    doc.add_paragraph(f"\nEducation:\n{education}")

    filename = f"CV_{index + 400}.docx"
    doc.save(os.path.join(output_dir, filename))

# Generate 200 CVs
for i in range(200):
    generate_cv(i)

print(f"✅ Generated 200 fake CVs in: {output_dir}")

# Move to CVs folder
dest_dir = "CVs"
os.makedirs(dest_dir, exist_ok=True)

for file in os.listdir(output_dir):
    shutil.copy(os.path.join(output_dir, file), os.path.join(dest_dir, file))

print(f"📥 Copied all fake CVs to {dest_dir} (your system's default folder)")
